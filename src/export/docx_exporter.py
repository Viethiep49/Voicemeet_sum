"""
Professional DOCX export for meeting minutes
File: src/export/docx_exporter.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from typing import Optional
from pathlib import Path

from ..utils.logger import logger


class MeetingDocxExporter:
    """Export meeting minutes to professional DOCX format"""

    def __init__(self):
        """Initialize DOCX exporter"""
        self.doc = None

    def export(
        self,
        extracted_data: dict,
        transcript: str,
        output_path: str,
        include_transcript: bool = True
    ) -> str:
        """
        Export meeting data to DOCX file

        Args:
            extracted_data: Structured meeting data from extractor
            transcript: Full transcript text
            output_path: Path to save DOCX file
            include_transcript: Whether to include full transcript as appendix

        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating DOCX export: {output_path}")

        # Create document
        self.doc = Document()
        self._setup_styles()

        # Add content sections
        self._add_header()
        self._add_meeting_info(extracted_data.get("meeting_info", {}))
        self._add_discussions(extracted_data.get("discussions", []))
        self._add_decisions(extracted_data.get("decisions", []))
        self._add_action_items(extracted_data.get("action_items", []))
        self._add_other_notes(extracted_data.get("other_notes"))

        if include_transcript:
            self._add_transcript(transcript)

        # Save document
        self.doc.save(output_path)
        logger.info(f"DOCX export completed: {output_path}")

        return output_path

    def _setup_styles(self):
        """Configure document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def _add_header(self):
        """Add document header/title"""
        # Title
        title = self.doc.add_heading('BIÊN BẢN CUỘC HỌP', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtitle with date
        date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        subtitle = self.doc.add_paragraph(f'Ngày tạo: {date_str}')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_format = subtitle.runs[0].font
        subtitle_format.size = Pt(10)
        subtitle_format.color.rgb = RGBColor(128, 128, 128)

        # Separator
        self.doc.add_paragraph('_' * 80)

    def _add_meeting_info(self, meeting_info: dict):
        """Add meeting information section"""
        self.doc.add_heading('1. THÔNG TIN CUỘC HỌP', level=1)

        # Main purpose
        main_purpose = meeting_info.get("main_purpose", "N/A")
        p = self.doc.add_paragraph()
        p.add_run('Mục đích: ').bold = True
        p.add_run(main_purpose)

        # Topics discussed
        topics = meeting_info.get("topics_discussed", [])
        if topics:
            p = self.doc.add_paragraph()
            p.add_run('Chủ đề thảo luận:').bold = True
            for topic in topics:
                self.doc.add_paragraph(topic, style='List Bullet')

        # Participants
        participants = meeting_info.get("participants_mentioned", [])
        if participants:
            p = self.doc.add_paragraph()
            p.add_run('Người tham gia: ').bold = True
            p.add_run(', '.join(participants) if participants else 'Không xác định')

        self.doc.add_paragraph()  # Spacing

    def _add_discussions(self, discussions: list):
        """Add discussions section"""
        if not discussions:
            return

        self.doc.add_heading('2. NỘI DUNG THẢO LUẬN', level=1)

        for idx, discussion in enumerate(discussions, 1):
            topic = discussion.get("topic", f"Chủ đề {idx}")

            # Topic heading
            self.doc.add_heading(f'2.{idx}. {topic}', level=2)

            # Discussion points
            points = discussion.get("points", [])
            for point in points:
                speaker = point.get("speaker")
                content = point.get("content", "")
                point_type = point.get("type", "opinion")

                # Format based on type
                p = self.doc.add_paragraph()

                if speaker:
                    run = p.add_run(f'{speaker}: ')
                    run.bold = True

                # Add type indicator
                type_indicators = {
                    "question": "❓ ",
                    "answer": "💬 ",
                    "decision": "✅ ",
                    "proposal": "💡 "
                }
                type_prefix = type_indicators.get(point_type, "• ")

                p.add_run(f'{type_prefix}{content}')
                p.style = 'List Bullet'

            # Conclusion
            conclusion = discussion.get("conclusion")
            if conclusion:
                p = self.doc.add_paragraph()
                p.add_run('Kết luận: ').bold = True
                p.add_run(conclusion)
                p_format = p.paragraph_format
                p_format.left_indent = Inches(0.5)

            self.doc.add_paragraph()  # Spacing

    def _add_decisions(self, decisions: list):
        """Add decisions section"""
        if not decisions:
            return

        self.doc.add_heading('3. CÁC QUYẾT ĐỊNH', level=1)

        for idx, decision in enumerate(decisions, 1):
            content = decision.get("content", "")
            made_by = decision.get("made_by")

            p = self.doc.add_paragraph()
            p.add_run(f'{idx}. ').bold = True
            p.add_run(content)

            if made_by:
                p.add_run(f' (Quyết định bởi: {made_by})')
                p.runs[-1].font.italic = True
                p.runs[-1].font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_paragraph()  # Spacing

    def _add_action_items(self, action_items: list):
        """Add action items section as a table"""
        if not action_items:
            return

        self.doc.add_heading('4. CÔNG VIỆC CẦN LÀM', level=1)

        # Create table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        headers = ['STT', 'Công việc', 'Người phụ trách', 'Deadline']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Bold header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add action items
        for idx, item in enumerate(action_items, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = item.get("task", "")
            row_cells[2].text = item.get("assignee") or "Chưa phân công"

            deadline = item.get("deadline") or "Chưa xác định"
            priority = item.get("priority")

            # Add priority indicator
            if priority == "high":
                deadline += " 🔴"
            elif priority == "medium":
                deadline += " 🟡"
            elif priority == "low":
                deadline += " 🟢"

            row_cells[3].text = deadline

        self.doc.add_paragraph()  # Spacing

    def _add_other_notes(self, notes: Optional[str]):
        """Add other notes section"""
        if not notes:
            return

        self.doc.add_heading('5. GHI CHÚ KHÁC', level=1)
        self.doc.add_paragraph(notes)
        self.doc.add_paragraph()  # Spacing

    def _add_transcript(self, transcript: str):
        """Add full transcript as appendix"""
        if not transcript:
            return

        # Page break
        self.doc.add_page_break()

        self.doc.add_heading('PHỤ LỤC: TRANSCRIPT ĐẦY ĐỦ', level=1)

        # Add transcript with smaller font
        p = self.doc.add_paragraph(transcript)
        p_format = p.paragraph_format
        p_format.line_spacing = 1.15

        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = 'Courier New'

    def _add_footer(self):
        """Add document footer (optional)"""
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = f"Generated by Voicemeet_sum - {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
